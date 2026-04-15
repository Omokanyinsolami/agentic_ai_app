from __future__ import annotations

import shutil
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "agentic ai dissertation.docx"
BACKUP_PATH = ROOT / "agentic ai dissertation.backup_before_live_edits.docx"
FIG_DIR = ROOT / "Latex" / "dissertation_figures"


def ensure_backup() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, box, text, *, fill, outline, title_font, body_font):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=20, fill=fill, outline=outline, width=3)
    lines = textwrap.wrap(text, width=20)
    current_y = y1 + 22
    for idx, line in enumerate(lines):
        font = title_font if idx == 0 and len(lines) > 1 else body_font
        bbox = draw.textbbox((0, 0), line, font=font)
        text_w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - text_w) / 2, current_y), line, font=font, fill="black")
        current_y += (bbox[3] - bbox[1]) + 8


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, *, fill="#334155", width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 8), (ex - 16 * direction, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - 16 * direction), (ex + 8, ey - 16 * direction)]
    draw.polygon(points, fill=fill)


def create_use_case_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    body_font = get_font(24)
    small_font = get_font(22)
    draw.text((80, 40), "Learner-Support Use Case Overview", font=title_font, fill="#0f172a")

    actor_box = (70, 300, 320, 520)
    draw.rounded_rectangle(actor_box, radius=28, fill="#dbeafe", outline="#2563eb", width=4)
    draw.ellipse((130, 330, 260, 460), fill="#93c5fd", outline="#2563eb", width=4)
    draw.line((195, 460, 195, 515), fill="#2563eb", width=5)
    draw.line((145, 520, 245, 520), fill="#2563eb", width=5)
    draw.line((195, 515, 145, 575), fill="#2563eb", width=5)
    draw.line((195, 515, 245, 575), fill="#2563eb", width=5)
    actor_text = "Learner"
    bbox = draw.textbbox((0, 0), actor_text, font=title_font)
    draw.text((195 - (bbox[2] - bbox[0]) / 2, 590), actor_text, font=title_font, fill="#0f172a")

    boxes = [
        ((430, 170, 800, 290), "Capture and manage tasks"),
        ((860, 170, 1230, 290), "Review tasks and workload"),
        ((430, 360, 800, 480), "Generate a study schedule"),
        ((860, 360, 1230, 480), "Adapt plans after change"),
        ((430, 550, 800, 670), "Understand priorities and conflicts"),
        ((860, 550, 1230, 670), "Receive reminders and support"),
    ]
    for box, label in boxes:
        draw_box(draw, box, label, fill="#ecfeff", outline="#0891b2", title_font=small_font, body_font=small_font)
        draw_arrow(draw, (320, 410), (box[0], (box[1] + box[3]) // 2))

    note = "Core learner-facing interactions carried into implementation, testing, and evaluation."
    draw.text((430, 760), note, font=body_font, fill="#334155")
    img.save(path)


def create_request_workflow_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "#fffdf7")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    body_font = get_font(24)
    draw.text((80, 40), "Operational Workflow for Requests, Scheduling, and Adaptation", font=title_font, fill="#111827")

    nodes = [
        ((90, 300, 320, 430), "Learner request"),
        ((390, 300, 620, 430), "React frontend"),
        ((690, 300, 920, 430), "Flask API"),
        ((990, 300, 1220, 430), "LangGraph router"),
        ((1290, 210, 1520, 340), "PostgreSQL"),
        ((1290, 470, 1520, 600), "Schedule / response"),
    ]
    colors = ["#dbeafe", "#fef3c7", "#dcfce7", "#ede9fe", "#fee2e2", "#cffafe"]
    outlines = ["#2563eb", "#d97706", "#16a34a", "#7c3aed", "#dc2626", "#0f766e"]
    for (box, label), fill, outline in zip(nodes, colors, outlines):
        draw_box(draw, box, label, fill=fill, outline=outline, title_font=body_font, body_font=body_font)

    draw_arrow(draw, (320, 365), (390, 365))
    draw_arrow(draw, (620, 365), (690, 365))
    draw_arrow(draw, (920, 365), (990, 365))
    draw_arrow(draw, (1220, 335), (1290, 275))
    draw_arrow(draw, (1220, 395), (1290, 535))
    draw_arrow(draw, (1290, 600), (390, 600))
    draw_arrow(draw, (390, 600), (390, 430))
    draw.text((450, 635), "validated output, persisted state, and user-facing explanations", font=body_font, fill="#374151")
    img.save(path)


def create_reminder_workflow_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    body_font = get_font(24)
    draw.text((80, 40), "Hosted Reminder Automation Workflow", font=title_font, fill="#111827")

    nodes = [
        ((80, 330, 310, 460), "Supabase Cron"),
        ((390, 330, 690, 460), "Edge Function:\ndeadline-reminders"),
        ((770, 180, 1070, 310), "Supabase\nPostgreSQL"),
        ((770, 480, 1070, 610), "Reminder log"),
        ((1150, 330, 1390, 460), "Brevo API"),
        ((1450, 330, 1570, 460), "Learner email"),
    ]
    fills = ["#fee2e2", "#ede9fe", "#dcfce7", "#fef9c3", "#dbeafe", "#cffafe"]
    outlines = ["#dc2626", "#7c3aed", "#16a34a", "#ca8a04", "#2563eb", "#0891b2"]
    for (box, label), fill, outline in zip(nodes, fills, outlines):
        draw_box(draw, box, label, fill=fill, outline=outline, title_font=body_font, body_font=body_font)

    draw_arrow(draw, (310, 395), (390, 395))
    draw_arrow(draw, (690, 360), (770, 245))
    draw_arrow(draw, (690, 430), (770, 545))
    draw_arrow(draw, (1070, 395), (1150, 395))
    draw_arrow(draw, (1390, 395), (1450, 395))
    draw.text((390, 700), "Triggers daily checks for due-soon and overdue pending tasks, then logs dispatch status.", font=body_font, fill="#334155")
    img.save(path)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_centered_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float = 6.4) -> Paragraph:
    pic_para = insert_paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return pic_para


def add_caption_after(paragraph: Paragraph, caption: str) -> Paragraph:
    caption_para = insert_paragraph_after(paragraph, caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_para.runs:
        caption_para.runs[0].font.size = Pt(10.5)
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.color.rgb = RGBColor(55, 65, 81)
    return caption_para


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for para in doc.paragraphs:
        if predicate(para.text.strip()):
            return para
    raise ValueError("Required paragraph not found")


def refresh_list_of_figures(doc: Document) -> None:
    heading = find_paragraph(doc, lambda t: t == "List of Figures")
    list_tables = find_paragraph(doc, lambda t: t == "List of Tables")

    to_remove = []
    seen = False
    for para in doc.paragraphs:
        if para._p == heading._p:
            seen = True
            continue
        if para._p == list_tables._p:
            break
        if seen:
            to_remove.append(para)

    for para in to_remove:
        remove_paragraph(para)

    entries = [
        "Figure 1. Use-case view of the agentic AI learner-support system",
        "Figure 2. Architecture of the agentic AI learner-support system for education",
        "Figure 3. Operational workflow for request handling, scheduling, and adaptation",
        "Figure 4. Hosted reminder automation workflow using Supabase and Brevo",
    ]

    current = heading
    for entry in entries:
        current = insert_paragraph_after(current, entry)


def update_references(doc: Document) -> None:
    replacements = {
        "Acharya, D. B.": "Acharya, D. B., Kuppan, K. & Divya, B. (2025) 'Agentic AI: Autonomous Intelligence for Complex Goals-A Comprehensive Survey', IEEE Access, 13, pp. 18912-18936. doi: 10.1109/ACCESS.2025.3532853.",
        "Bandi, A.": "Bandi, A., Kongari, B., Naguru, R., Pasnoor, S. & Vilipala, S. V. (2025) 'The rise of agentic AI: A review of definitions, frameworks, architectures, applications, evaluation metrics, and challenges', Future Internet, 17(9), article 404.",
        "Chase, H.": "Chase, H. (2022) LangChain [online]. Available at: https://github.com/langchain-ai/langchain (Accessed: 2026-04-08).",
        "CrewAI (2024)": "CrewAI (2024) CrewAI [online]. Available at: https://github.com/crewAIInc/crewAI (Accessed: 2026-04-08).",
        "Hughes, L.": "Hughes, L. et al. (2025) 'AI agents and agentic systems: A multi-expert analysis', Journal of Computer Information Systems, pp. 1-29.",
        "Intelligence, A.": "Russell, S. & Norvig, P. (2002) Artificial Intelligence: A Modern Approach. Upper Saddle River, NJ: Prentice Hall.",
        "LangChain AI (2024)": "LangChain AI (2024) LangGraph overview [online]. Available at: https://docs.langchain.com/oss/python/langgraph/overview (Accessed: 2026-04-08).",
        "Lee, C. P.": "Lee, C. P., Porfirio, D., Wang, X. J., Zhao, K. C. & Mutlu, B. (2025) 'VeriPlan: Integrating formal verification and LLMs into end-user planning', Proceedings of the 2025 CHI Conference on Human Factors in Computing Systems, pp. 1-19.",
        "Meta Platforms (2013)": "Meta Platforms (2013) React [online]. Available at: https://react.dev/ (Accessed: 2026-04-08).",
        "Ronacher, A.": "Ronacher, A. (2010) Flask [online]. Available at: https://flask.palletsprojects.com/ (Accessed: 2026-04-08).",
        "Shavit, Y.": "Shavit, Y., Agarwal, S., Brundage, M. et al. (2023) Practices for Governing Agentic AI Systems [online]. OpenAI. Available at: https://openai.com/index/practices-for-governing-agentic-ai-systems/ (Accessed: 2026-04-08).",
        "Sun, H.": "Sun, H., Zhuang, Y., Kong, L., Dai, B. & Zhang, C. (2023) 'AdaPlanner: Adaptive Planning from Feedback with Language Models', Advances in Neural Information Processing Systems, 36, pp. 58202-58245.",
        "The PostgreSQL Global Development Group": "The PostgreSQL Global Development Group (1996) PostgreSQL [online]. Available at: https://www.postgresql.org/ (Accessed: 2026-04-08).",
        "Wu, Q.": "Wu, Q. et al. (2023) 'AutoGen: Enabling Next-Generation LLM Applications via Multi-Agent Conversation' [online]. Available at: https://github.com/microsoft/autogen (Accessed: 2026-04-08).",
        "Yao, S.": "Yao, S. et al. (2022) 'ReAct: Synergizing Reasoning and Acting in Language Models', The Eleventh International Conference on Learning Representations.",
    }

    in_refs = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "References":
            in_refs = True
            continue
        if text == "Appendix":
            break
        if not in_refs or not text:
            continue
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                para.text = replacement
                break


def add_figures(doc: Document) -> None:
    use_cases_anchor = find_paragraph(doc, lambda t: t.startswith("These use cases were not treated as narrative examples only."))
    intro2 = insert_paragraph_after(
        use_cases_anchor,
        "Figure 1 summarises the main learner-facing interactions that the implemented system was designed to support across capture, scheduling, adaptation, explanation, and reminders."
    )
    pic2 = add_centered_picture_after(intro2, FIG_DIR / "figure_use_case_overview.png")
    add_caption_after(pic2, "Figure 1. Use-case view of the agentic AI learner-support system")

    arch_text = find_paragraph(doc, lambda t: t.startswith("The implemented system follows a three-tier architecture"))
    arch_text.text = arch_text.text.replace("Figure 1", "Figure 2")
    arch_caption = find_paragraph(doc, lambda t: t.startswith("Figure 1. Prototype architecture"))
    arch_caption.text = "Figure 2. Architecture of the agentic AI learner-support system for education"

    workflow_intro = insert_paragraph_after(
        arch_caption,
        "Figure 3 complements the architecture view by showing how user requests move through the frontend, backend, LangGraph routing layer, persistence layer, and response path during scheduling and adaptation."
    )
    pic3 = add_centered_picture_after(workflow_intro, FIG_DIR / "figure_request_workflow.png")
    add_caption_after(pic3, "Figure 3. Operational workflow for request handling, scheduling, and adaptation")

    reminder_anchor = find_paragraph(doc, lambda t: t.startswith("An additional setup outcome was the preparation of an always-on reminder route."))
    reminder_intro = insert_paragraph_after(
        reminder_anchor,
        "Figure 4 shows the hosted reminder path used for daily due-soon and overdue notifications when the Supabase and Brevo configuration is enabled."
    )
    pic4 = add_centered_picture_after(reminder_intro, FIG_DIR / "figure_reminder_workflow.png")
    add_caption_after(pic4, "Figure 4. Hosted reminder automation workflow using Supabase and Brevo")


def main() -> None:
    ensure_backup()
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    create_use_case_diagram(FIG_DIR / "figure_use_case_overview.png")
    create_request_workflow_diagram(FIG_DIR / "figure_request_workflow.png")
    create_reminder_workflow_diagram(FIG_DIR / "figure_reminder_workflow.png")

    doc = Document(str(DOCX_PATH))
    refresh_list_of_figures(doc)
    add_figures(doc)
    update_references(doc)
    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    main()
