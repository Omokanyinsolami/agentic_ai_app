from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


DOC_PATH = Path("agentic ai dissertation.docx")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def backup(path: Path) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    target = path.with_name(f"{path.stem}.backup_before_strict_reference_cleanup_{stamp}{path.suffix}")
    shutil.copy2(path, target)
    return target


def normalize_body_citations(text: str) -> list[str]:
    matches = re.findall(r"\(([^)]*?(?:19|20)\d{2}[^)]*?)\)", text)
    return [m.strip() for m in matches]


def main() -> None:
    backup_path = backup(DOC_PATH)
    doc = Document(DOC_PATH)

    replacements = {
        "The project is motivated by a practical support problem and a technical design problem.": (
            "The project is motivated by a practical support problem and a technical design problem. "
            "The practical problem is that many existing educational support tools are passive: they record deadlines or display events but expect learners to do most of the prioritisation and replanning work themselves. "
            "The technical problem is how to design an AI system that supports users actively without becoming opaque, unreliable, or difficult to assess."
        ),
        "Three specific motivations follow from this.": (
            "Three specific motivations follow from this. "
            "First, there is a need for a supportive environment in which learners can capture academic tasks, understand what matters next, and reorganise plans when constraints change. "
            "Second, there is a need to compare agentic support with existing passive or non-agentic solutions, because a system should justify why active AI assistance is useful rather than simply novel. "
            "Third, there is a need for an implementation that remains operational and defensible: setup, configuration, testing, and evaluation must show that the system works as a software prototype and not merely as a conceptual design."
        ),
        "What is often missing is a coherent supportive environment that actively combines these functions.": (
            "What is often missing is a coherent supportive environment that actively combines these functions. "
            "Existing systems rarely integrate task capture, prioritisation, scheduling, conflict detection, explanation, and adaptation into one connected workflow. "
            "This creates a gap between digital access to academic information and genuine support for academic self-management."
        ),
        "This project responds to that gap by treating learner support as an operational design objective.": (
            "This project responds to that gap by treating learner support as an operational design objective. "
            "The prototype is therefore intended to help learners manage academic tasks more effectively, not merely provide another storage interface. "
            "To justify an agentic approach, the system must do more than record tasks: it must actively prioritize, schedule, explain, detect problems, and adapt when circumstances change."
        ),
        "In addition to conceptual literature, the project required a technical review of practical architecture patterns for reliable agent-enabled software systems.": (
            "In addition to conceptual literature, the project required a technical review of practical architecture patterns for reliable agent-enabled software systems. "
            "This review focused on a central engineering question: how can an application remain predictable and testable while still benefiting from adaptive AI behaviour? "
            "The concern is not only model quality; it is also control flow, state integrity, and failure recovery under realistic operational conditions."
        ),
        "This project is framed as a design-and-implementation study of a functional learner-support prototype that uses agentic AI techniques to assist with academic task management.": (
            "This project is framed as a design-and-implementation study of a functional learner-support prototype that uses agentic AI techniques to assist with academic task management. "
            "It is therefore both a software prototype project and an evaluative AI systems project. "
            "The design focus is on building a usable support environment for learners, while the research focus is on understanding whether agentic orchestration improves usefulness, adaptability, and operational behaviour compared with passive or non-agentic alternatives."
        ),
    }

    # Precise replacements for paragraphs that previously had coverage-driven tail sentences.
    exact_replacements = {
        "The project is motivated by a practical support problem and a technical design problem. The practical problem is that many existing educational support tools are passive: they record deadlines or display events but expect learners to do most of the prioritisation and replanning work themselves. The technical problem is how to design an AI system that supports users actively without becoming opaque, unreliable, or difficult to assess. This concern is consistent with wider discussions of agentic AI adoption, where capability must be balanced against reliability, governance, and user trust (Murugesan, 2025).":
            replacements["The project is motivated by a practical support problem and a technical design problem."],
        "Three specific motivations follow from this. First, there is a need for a supportive environment in which learners can capture academic tasks, understand what matters next, and reorganise plans when constraints change. Second, there is a need to compare agentic support with existing passive or non-agentic solutions, because a system should justify why active AI assistance is useful rather than simply novel. Third, there is a need for an implementation that remains operational and defensible: setup, configuration, testing, and evaluation must show that the system works as a software prototype and not merely as a conceptual design. The shift from generative AI toward agentic AI also places stronger emphasis on action, feedback, and operational accountability rather than text generation alone (Schneider, 2025).":
            replacements["Three specific motivations follow from this."],
        "What is often missing is a coherent supportive environment that actively combines these functions. Existing systems rarely integrate task capture, prioritisation, scheduling, conflict detection, explanation, and adaptation into one connected workflow. This creates a gap between digital access to academic information and genuine support for academic self-management. Automated planning research in dynamic domains also highlights the importance of connecting plan generation with changing situational constraints (Bidoux et al., 2017).":
            replacements["What is often missing is a coherent supportive environment that actively combines these functions."],
        "This project responds to that gap by treating learner support as an operational design objective. The prototype is therefore intended to help learners manage academic tasks more effectively, not merely provide another storage interface. To justify an agentic approach, the system must do more than record tasks: it must actively prioritize, schedule, explain, detect problems, and adapt when circumstances change. This is consistent with recent agentic AI surveys that position autonomous support as an applied systems problem rather than only a modelling problem (Pati, 2025).":
            replacements["This project responds to that gap by treating learner support as an operational design objective."],
        "In addition to conceptual literature, the project required a technical review of practical architecture patterns for reliable agent-enabled software systems. This review focused on a central engineering question: how can an application remain predictable and testable while still benefiting from adaptive AI behaviour? The concern is not only model quality; it is also control flow, state integrity, and failure recovery under realistic operational conditions. Research on self-evolving agents reinforces the need to examine architecture, feedback loops, and control boundaries rather than treating autonomy as a single feature (Fang et al., 2025).":
            replacements["In addition to conceptual literature, the project required a technical review of practical architecture patterns for reliable agent-enabled software systems."],
        "This project is framed as a design-and-implementation study of a functional learner-support prototype that uses agentic AI techniques to assist with academic task management. It is therefore both a software prototype project and an evaluative AI systems project. The design focus is on building a usable support environment for learners, while the research focus is on understanding whether agentic orchestration improves usefulness, adaptability, and operational behaviour compared with passive or non-agentic alternatives. Recent research landscapes of agentic AI and LLMs also emphasise that evaluation should consider applications, limitations, and deployment implications together (Brohi et al., 2025).":
            replacements["This project is framed as a design-and-implementation study of a functional learner-support prototype that uses agentic AI techniques to assist with academic task management."],
        "CrewAI (2024) CrewAI [online]. Available at: https://github.com/crewAIInc/crewAI (Accessed: 9 April 2026).":
            "CrewAI uses a role-based coordination model, but this pattern aligned less well with the structured task-management workflow required in this dissertation (CrewAI, 2024).",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in exact_replacements:
            paragraph.text = exact_replacements[text]

    paragraphs = [p.text.strip() for p in doc.paragraphs]
    try:
        ref_index = paragraphs.index("References")
    except ValueError:
        ref_index = None

    if ref_index is None:
        raise RuntimeError("Could not find References heading in dissertation document.")

    body_text = "\n".join(t for t in paragraphs[:ref_index] if t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        body_text += "\n" + text

    cited = set(normalize_body_citations(body_text))

    ref_map = {
        "Acharya, D. B., Kuppan, K. and Divya, B. (2025)": "Acharya et al., 2025",
        "Aghzal, M., Plaku, E., Stein, G. J. and Yao, Z. (2025)": "Aghzal et al., 2025",
        "Bandi, A., Kongari, B., Naguru, R., Pasnoor, S. and Vilipala, S. V. (2025)": "Bandi et al., 2025",
        "Bidoux, L., Pignon, J. P. and Bénaben, F. (2017)": "Bidoux et al., 2017",
        "Brohi, S., Jhanjhi, N. and Pillai, T. R. (2025)": "Brohi et al., 2025",
        "Chase, H. (2022)": "Chase, 2022",
        "Chen, Z., Tang, J., Chen, X. and Lin, Y. (2023)": "Chen et al., 2023",
        "CrewAI (2024)": "CrewAI, 2024",
        "Fang, J. et al. (2025)": "Fang et al., 2025",
        "Geffner, H. and Bonet, B. (2013)": "Geffner and Bonet, 2013",
        "Ghallab, M., Nau, D. and Traverso, P. (2016)": "Ghallab et al., 2016",
        "Hughes, L. et al. (2025)": "Hughes et al., 2025",
        "Russell, S. and Norvig, P. (2021)": "Russell and Norvig, 2021",
        "Katz, M., Kokel, H., Srinivas, K. and Sohrabi Araghi, S. (2024)": "Katz et al., 2024",
        "Koedinger, K. R., Baker, R. S., Cunningham, K., Skogsholm, A., Leber, B. and Stamper, J. (2010)": "Koedinger et al., 2010",
        "LangChain AI (2024)": "LangChain AI, 2024",
        "Lee, C. P., Porfirio, D., Wang, X. J., Zhao, K. C. and Mutlu, B. (2025)": "Lee et al., 2025",
        "Meta Platforms (2026)": "Meta Platforms, 2026",
        "Murugesan, S. (2025)": "Murugesan, 2025",
        "Novák, P. and Jamroga, W. (2011)": "Novák and Jamroga, 2011",
        "Pati, A. K. (2025)": "Pati, 2025",
        "Ronacher, A. (2010)": "Ronacher, 2010",
        "Sapkota, R., Roumeliotis, K. I. and Karkee, M. (2025)": "Sapkota et al., 2025",
        "Schneider, J. (2025)": "Schneider, 2025",
        "Shavit, Y., Agarwal, S., Brundage, M. et al. (2023)": "Shavit et al., 2023",
        "Sun, H., Zhuang, Y., Kong, L., Dai, B. and Zhang, C. (2023)": "Sun et al., 2023",
        "The PostgreSQL Global Development Group (2026)": "The PostgreSQL Global Development Group, 2026",
        "Wang, L. et al. (2024)": "Wang et al., 2024",
        "Wu, Q. et al. (2023)": "Wu et al., 2023",
        "Yao, S. et al. (2022)": "Yao et al., 2022",
    }

    uncited_removed = []
    found_references = False
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "References":
            found_references = True
            continue
        if not found_references:
            continue
        if not text:
            continue
        if text.startswith("Appendix"):
            break
        matched_key = None
        for prefix, citation_key in ref_map.items():
            if text.startswith(prefix):
                matched_key = citation_key
                break
        if matched_key and matched_key not in cited:
            uncited_removed.append(matched_key)
            remove_paragraph(paragraph)

    doc.save(DOC_PATH)

    refreshed = Document(DOC_PATH)
    paras = [p.text.strip() for p in refreshed.paragraphs]
    ref_index = paras.index("References")
    body_text = "\n".join(t for t in paras[:ref_index] if t)
    for table in refreshed.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        body_text += "\n" + text

    body_citations = set(normalize_body_citations(body_text))
    refs_remaining = []
    for text in paras[ref_index + 1 :]:
        if not text:
            continue
        if text.startswith("Appendix"):
            break
        if re.match(r"^\d+(\.\d+)?\s+", text):
            break
        for prefix, citation_key in ref_map.items():
            if text.startswith(prefix):
                refs_remaining.append(citation_key)
                break

    print(f"backup={backup_path}")
    print(f"uncited_removed={uncited_removed}")
    print(f"remaining_reference_count={len(refs_remaining)}")
    print(f"remaining_references={refs_remaining}")
    print(f"body_citation_count={len(body_citations)}")


if __name__ == "__main__":
    main()
