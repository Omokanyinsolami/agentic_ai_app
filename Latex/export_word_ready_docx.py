import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parent
MAIN_TEX = ROOT / "main.tex"
REF_BIB = ROOT / "references.bib"
OUTPUT = ROOT / "main_word_ready.docx"
GOOGLE_OUTPUT = ROOT / "main_google_docs_ready.docx"


SIMPLE_TEXT_REPLACEMENTS = [
    ("~", " "),
    ("\\_", "_"),
    ("\\&", "&"),
    ("\\%", "%"),
    ("\\#", "#"),
    ("\\$", "$"),
    ("\\{", "{"),
    ("\\}", "}"),
    ("---", "-"),
    ("--", "-"),
    ("``", '"'),
    ("''", '"'),
]

LATEX_SYMBOL_REPLACEMENTS = {
    r"\ae": "ae",
    r"\AE": "AE",
    r"\oe": "oe",
    r"\OE": "OE",
    r"\aa": "aa",
    r"\AA": "AA",
    r"\o": "o",
    r"\O": "O",
    r"\ss": "ss",
    r"\l": "l",
    r"\L": "L",
}

ACCENT_MARKS = {
    "'": "\u0301",
    "`": "\u0300",
    "^": "\u0302",
    '"': "\u0308",
    "~": "\u0303",
    "=": "\u0304",
    ".": "\u0307",
    "u": "\u0306",
    "v": "\u030C",
    "H": "\u030B",
    "c": "\u0327",
    "k": "\u0328",
    "r": "\u030A",
    "b": "\u0331",
}


@dataclass
class TableBuffer:
    caption: str = ""
    rows: list[list[str]] = field(default_factory=list)
    raw_row_buffer: str = ""


@dataclass
class FigureBuffer:
    caption: str = ""
    image_path: Path | None = None


@dataclass
class OutlineEntry:
    level: int
    title: str


@dataclass
class HeadingCounters:
    section: int = 0
    subsection: int = 0
    subsubsection: int = 0


@dataclass
class BibEntry:
    key: str
    entry_type: str
    raw_author: str
    year: str
    title: str
    journal: str
    booktitle: str
    publisher: str
    howpublished: str
    note: str
    volume: str
    number: str
    pages: str
    doi: str


def replace_commands_with_arg(text: str, commands: list[str]) -> str:
    for command in commands:
        text = re.sub(rf"\\{command}\*?\{{([^{{}}]*)\}}", r"\1", text)
    return text


def apply_accent_macros(text: str) -> str:
    def repl(match):
        accent = match.group(1)
        base = match.group(2)
        combined = unicodedata.normalize("NFC", base + ACCENT_MARKS[accent])
        return combined

    pattern = r"\\([\'`\^\"~=\.uvHckrb])\{?([A-Za-z])\}?"
    return re.sub(pattern, repl, text)


def latex_to_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\n", " ")
    text = apply_accent_macros(text)
    for old, new in SIMPLE_TEXT_REPLACEMENTS:
        text = text.replace(old, new)
    for old, new in LATEX_SYMBOL_REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def clean_inline(text: str) -> str:
    text = text.strip()
    if not text:
        return ""

    text = re.sub(r"(?<!\\)%.*$", "", text).strip()
    text = replace_citation_commands(text)
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\ref\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\addcontentsline\{[^}]*\}\{[^}]*\}\{[^}]*\}", "", text)
    text = re.sub(r"\\newline", "", text)
    text = re.sub(r"\\captionof\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = replace_commands_with_arg(
        text,
        ["texttt", "textbf", "emph", "textit", "underline", "textsc", "texorpdfstring"],
    )
    text = re.sub(r"\\multicolumn\{[^}]*\}\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\multirow\{[^}]*\}\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\cellcolor\{[^}]*\}", "", text)
    text = text.replace("\\ganttcell", "X")
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = latex_to_text(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_macro(tex: str, macro: str, default: str = "") -> str:
    match = re.search(rf"\\{macro}\{{(.*?)\}}", tex, re.S)
    if not match:
        return default
    return clean_inline(match.group(1))


def parse_braced_value(text: str, start_index: int) -> tuple[str, int]:
    depth = 0
    chars = []
    index = start_index
    while index < len(text):
        char = text[index]
        if char == "{":
            if depth > 0:
                chars.append(char)
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return "".join(chars), index + 1
            chars.append(char)
        else:
            chars.append(char)
        index += 1
    return "".join(chars), index


def parse_quoted_value(text: str, start_index: int) -> tuple[str, int]:
    chars = []
    index = start_index + 1
    while index < len(text):
        char = text[index]
        if char == '"' and text[index - 1] != "\\":
            return "".join(chars), index + 1
        chars.append(char)
        index += 1
    return "".join(chars), index


def extract_bib_field(block: str, field_name: str) -> str:
    match = re.search(rf"\b{field_name}\s*=\s*", block, re.I)
    if not match:
        return ""

    index = match.end()
    while index < len(block) and block[index].isspace():
        index += 1
    if index >= len(block):
        return ""

    if block[index] == "{":
        value, _ = parse_braced_value(block, index)
        return value
    if block[index] == '"':
        value, _ = parse_quoted_value(block, index)
        return value

    end_index = block.find(",", index)
    if end_index == -1:
        end_index = len(block)
    return block[index:end_index].strip()


def clean_bib_field(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = replace_commands_with_arg(
        text,
        ["texttt", "textbf", "emph", "textit", "underline", "textsc", "texorpdfstring"],
    )
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    return latex_to_text(text)


def parse_bib_database(bib_path: Path) -> dict[str, BibEntry]:
    if not bib_path.exists():
        return {}

    data = bib_path.read_text(encoding="utf-8", errors="ignore")
    entries: dict[str, BibEntry] = {}
    for block in re.split(r"\n@", "\n" + data):
        block = block.strip()
        if not block:
            continue
        if not block.startswith("@"):
            block = "@" + block

        entry_match = re.match(r"@(\w+)\s*\{\s*([^,]+),", block, re.S)
        if not entry_match:
            continue

        key = latex_to_text(entry_match.group(2).strip())
        entries[key] = BibEntry(
            key=key,
            entry_type=entry_match.group(1).lower(),
            raw_author=extract_bib_field(block, "author"),
            year=clean_bib_field(extract_bib_field(block, "year")),
            title=clean_bib_field(extract_bib_field(block, "title")),
            journal=clean_bib_field(extract_bib_field(block, "journal")),
            booktitle=clean_bib_field(extract_bib_field(block, "booktitle")),
            publisher=clean_bib_field(extract_bib_field(block, "publisher")),
            howpublished=clean_bib_field(extract_bib_field(block, "howpublished")),
            note=clean_bib_field(extract_bib_field(block, "note")),
            volume=clean_bib_field(extract_bib_field(block, "volume")),
            number=clean_bib_field(extract_bib_field(block, "number")),
            pages=clean_bib_field(extract_bib_field(block, "pages")),
            doi=clean_bib_field(extract_bib_field(block, "doi")),
        )
    return entries


def split_bib_authors(raw_author: str) -> list[str]:
    if not raw_author:
        return []

    authors = []
    for part in re.split(r"\s+and\s+", raw_author):
        part = part.strip()
        if not part or part.lower() == "others":
            continue
        authors.append(part)
    return authors


def is_corporate_author(raw_author: str) -> bool:
    stripped = raw_author.strip()
    return stripped.startswith("{") and stripped.endswith("}")


def clean_author_name(raw_author: str) -> str:
    return latex_to_text(raw_author.strip().strip("{}"))


def author_surname(raw_author: str) -> str:
    if not raw_author:
        return "Unknown"
    if is_corporate_author(raw_author):
        return clean_author_name(raw_author)

    clean = clean_author_name(raw_author)
    if "," in clean:
        return clean.split(",", 1)[0].strip()

    tokens = clean.split()
    return tokens[-1] if tokens else clean


def author_initials(raw_author: str) -> str:
    clean = clean_author_name(raw_author)
    if "," in clean:
        given_names = clean.split(",", 1)[1].strip()
    else:
        parts = clean.split()
        given_names = " ".join(parts[:-1])

    initials = []
    for token in re.split(r"[\s\-]+", given_names):
        token = token.strip()
        if token:
            initials.append(f"{token[0].upper()}.")
    return " ".join(initials)


def format_author_for_reference(raw_author: str) -> str:
    if not raw_author:
        return "Unknown"
    if is_corporate_author(raw_author):
        return clean_author_name(raw_author)

    surname = author_surname(raw_author)
    initials = author_initials(raw_author)
    return f"{surname}, {initials}".strip().rstrip(",")


def format_author_list_for_reference(raw_author: str) -> str:
    authors = split_bib_authors(raw_author)
    if not authors:
        return "Unknown"

    formatted = [format_author_for_reference(author) for author in authors]
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]} & {formatted[1]}"
    if len(formatted) <= 6:
        return ", ".join(formatted[:-1]) + f" & {formatted[-1]}"
    return f"{formatted[0]} et al."


def format_author_list_for_citation(raw_author: str) -> str:
    authors = split_bib_authors(raw_author)
    if not authors:
        return "Unknown"

    surnames = [author_surname(author) for author in authors]
    if len(surnames) == 1:
        return surnames[0]
    if len(surnames) == 2:
        return f"{surnames[0]} & {surnames[1]}"
    return f"{surnames[0]} et al."


def format_pages(pages: str) -> str:
    if not pages:
        return ""
    return f"pp. {pages.replace('--', '–').replace('-', '–')}"


def format_reference_entry(entry: BibEntry) -> str:
    author = format_author_list_for_reference(entry.raw_author)
    year = entry.year or "n.d."
    title = f"'{entry.title}'" if entry.title else ""
    source = entry.journal or entry.booktitle
    vol_issue = ""
    if entry.volume:
        vol_issue = entry.volume
        if entry.number:
            vol_issue += f"({entry.number})"
    pages = format_pages(entry.pages)

    parts = [f"{author} ({year})"]
    if title:
        parts.append(title)
    if source:
        parts.append(source)
    if vol_issue:
        parts.append(vol_issue)
    if pages:
        parts.append(pages)
    if entry.publisher and entry.publisher not in parts:
        parts.append(entry.publisher)
    if entry.howpublished:
        parts.append(f"Available at: {entry.howpublished}")
    if entry.note:
        parts.append(f"({entry.note})")
    if entry.doi:
        parts.append(f"doi:{entry.doi}")

    return ", ".join(part for part in parts if part).rstrip(".") + "."


def parse_bib_entries(bib_path: Path) -> list[str]:
    database = parse_bib_database(bib_path)
    entries = sorted(
        database.values(),
        key=lambda entry: (
            format_author_list_for_citation(entry.raw_author).lower(),
            (entry.year or "").lower(),
            (entry.title or "").lower(),
        ),
    )
    return [format_reference_entry(entry) for entry in entries]


BIB_DATABASE = parse_bib_database(REF_BIB)


def format_citation(keys_text: str, textcite: bool = False) -> str:
    citations = []
    for raw_key in keys_text.split(","):
        key = latex_to_text(raw_key.strip())
        entry = BIB_DATABASE.get(key)
        if entry:
            citations.append((format_author_list_for_citation(entry.raw_author), entry.year or "n.d."))
        elif key:
            citations.append((key, "n.d."))

    if not citations:
        return ""

    if textcite and len(citations) == 1:
        author, year = citations[0]
        return f"{author} ({year})"

    return "(" + "; ".join(f"{author}, {year}" for author, year in citations) + ")"


def replace_citation_commands(text: str) -> str:
    text = re.sub(r"\\textcite\{([^}]*)\}", lambda match: format_citation(match.group(1), textcite=True), text)
    text = re.sub(r"\\parencite\{([^}]*)\}", lambda match: format_citation(match.group(1)), text)
    text = re.sub(r"\\cite\{([^}]*)\}", lambda match: format_citation(match.group(1)), text)
    return text


def collect_caption_metadata(lines: list[str]) -> tuple[list[str], list[str]]:
    table_captions: list[str] = []
    figure_captions: list[str] = []
    table_depth = 0
    figure_depth = 0

    for raw in lines:
        line = raw.strip()
        if r"\begin{table}" in line:
            table_depth += 1
        if r"\begin{figure}" in line:
            figure_depth += 1

        match = re.search(r"\\caption\{(.+?)\}", line)
        if match:
            caption = clean_inline(match.group(1))
            if table_depth > 0:
                table_captions.append(caption)
            elif figure_depth > 0:
                figure_captions.append(caption)

        if r"\end{table}" in line and table_depth > 0:
            table_depth -= 1
        if r"\end{figure}" in line and figure_depth > 0:
            figure_depth -= 1

    return table_captions, figure_captions


def collect_outline_metadata(lines: list[str]) -> list[OutlineEntry]:
    outline: list[OutlineEntry] = []
    in_document = False
    in_titlepage = False
    numbering_active = False
    counters = HeadingCounters()

    for raw in lines:
        line = raw.strip()

        if r"\begin{document}" in line:
            in_document = True
            continue
        if not in_document:
            continue
        if r"\end{document}" in line:
            break

        if r"\begin{titlepage}" in line:
            in_titlepage = True
            continue
        if r"\end{titlepage}" in line:
            in_titlepage = False
            continue
        if in_titlepage:
            continue

        for pattern, level in (
            (r"\\section(\*?)\{(.+?)\}", 1),
            (r"\\subsection(\*?)\{(.+?)\}", 2),
            (r"\\subsubsection(\*?)\{(.+?)\}", 3),
        ):
            match = re.match(pattern, line)
            if match:
                starred = bool(match.group(1))
                title = clean_inline(match.group(2))
                if title and title not in {"Table of Contents", "List of Figures", "List of Tables"}:
                    display_title = title
                    if not starred:
                        if level == 1:
                            counters.section += 1
                            counters.subsection = 0
                            counters.subsubsection = 0
                            numbering_active = True
                            display_title = f"{counters.section} {title}"
                        elif level == 2 and numbering_active:
                            counters.subsection += 1
                            counters.subsubsection = 0
                            display_title = f"{counters.section}.{counters.subsection} {title}"
                        elif level == 3 and numbering_active:
                            counters.subsubsection += 1
                            display_title = f"{counters.section}.{counters.subsection}.{counters.subsubsection} {title}"
                    outline.append(OutlineEntry(level=level, title=display_title))
                break

    return outline


def add_field(paragraph, instruction: str, placeholder: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instruction_text)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_section_page_numbering(section, start: int | None = None, fmt: str | None = None):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)

    start_attr = qn("w:start")
    fmt_attr = qn("w:fmt")

    if start is None:
        pg_num_type.attrib.pop(start_attr, None)
    else:
        pg_num_type.set(start_attr, str(start))

    if fmt is None:
        pg_num_type.attrib.pop(fmt_attr, None)
    else:
        pg_num_type.set(fmt_attr, fmt)


def add_page_number_footer(section, instruction: str = "PAGE", placeholder: str = "1"):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, instruction, placeholder)


def apply_page_numbering(document: Document):
    sections = document.sections
    if not sections:
        return

    title_section = sections[0]
    title_section.footer.is_linked_to_previous = False
    for paragraph in title_section.footer.paragraphs:
        clear_paragraph(paragraph)

    if len(sections) >= 2:
        front_matter = sections[1]
        set_section_page_numbering(front_matter, start=1, fmt="lowerRoman")
        add_page_number_footer(front_matter, instruction=r"PAGE \* roman", placeholder="i")

    if len(sections) >= 3:
        body = sections[2]
        set_section_page_numbering(body, start=1, fmt="decimal")
        add_page_number_footer(body)

    for section in sections[3:]:
        add_page_number_footer(section)


def add_front_matter(document: Document):
    section = document.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    enable_update_fields_on_open(document)


def enable_update_fields_on_open(document: Document):
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def add_toc(document: Document):
    paragraph = document.add_paragraph()
    add_field(paragraph, r'TOC \o "1-3" \h \z \u', "Update table of contents in Word (Ctrl+A, then F9).")


def add_static_toc(document: Document, outline: list[OutlineEntry]):
    if not outline:
        document.add_paragraph("Table of contents entries were not detected during export.")
        return

    for entry in outline:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(6 * (entry.level - 1))
        paragraph.add_run(entry.title)


def add_caption_index(document: Document, label: str):
    paragraph = document.add_paragraph()
    add_field(
        paragraph,
        rf'TOC \h \z \c "{label}"',
        f"Update list of {label.lower()}s in Word (Ctrl+A, then F9).",
    )


def add_caption_paragraph(document: Document, label: str, caption_text: str):
    if not caption_text:
        return

    style = "Caption" if "Caption" in document.styles else "Normal"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(f"{label} ")
    add_field(paragraph, f"SEQ {label} \\* ARABIC", "1")
    paragraph.add_run(f". {caption_text}")


def add_static_caption_list(document: Document, label: str, captions: list[str]):
    if not captions:
        document.add_paragraph(f"No {label.lower()} captions were detected during export.")
        return

    for index, caption in enumerate(captions, start=1):
        document.add_paragraph(f"{label} {index}. {caption}")


def add_static_numbered_item(document: Document, depth: int, number: int, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Mm(8 * max(depth - 1, 0))
    # Google Docs aggressively merges plain "1. " paragraphs into one long list.
    # Insert a zero-width space after the period so the visible numbering remains
    # the same but importers treat each paragraph as literal text, not auto-listing.
    paragraph.add_run(f"{number}.\u200b {text}")


def flush_pending_paragraph(document: Document, pending_paragraph_lines: list[str]):
    if not pending_paragraph_lines:
        return []

    text = clean_inline(" ".join(pending_paragraph_lines))
    if text:
        document.add_paragraph(text)
    return []


def add_table_row(buffer: TableBuffer, row_text: str):
    row_text = row_text.replace(r"\hline", " ")
    row_text = re.sub(r"\\cline\{[^}]*\}", " ", row_text)
    row_text = row_text.strip()
    if not row_text:
        return

    cells = [clean_table_cell(cell) for cell in row_text.split("&")]
    if not any(cell for cell in cells):
        return
    buffer.rows.append(cells)


def clean_table_cell(text: str) -> str:
    text = text.strip()
    if not text:
        return ""

    multicolumn = re.search(r"\\multicolumn\{[^}]*\}\{[^}]*\}\{(.+)\}", text)
    if multicolumn:
        text = multicolumn.group(1)

    text = text.replace(r"\ganttcell", "X")
    text = re.sub(r"\\cellcolor\{[^}]*\}", "", text)
    return clean_inline(text)


def append_table_source(buffer: TableBuffer, line: str):
    stripped = re.sub(r"(?<!\\)%.*$", "", line).strip()
    if not stripped:
        return
    if stripped in {r"\centering", r"\small"}:
        return
    if stripped.startswith(r"\setlength") or stripped.startswith(r"\renewcommand") or stripped.startswith(r"\newcommand"):
        return
    if stripped.startswith(r"\label"):
        return

    cleaned = stripped.replace(r"\hline", " ")
    cleaned = re.sub(r"\\cline\{[^}]*\}", " ", cleaned)
    buffer.raw_row_buffer = f"{buffer.raw_row_buffer} {cleaned}".strip()

    while r"\\" in buffer.raw_row_buffer:
        row_text, remainder = buffer.raw_row_buffer.split(r"\\", 1)
        add_table_row(buffer, row_text)
        buffer.raw_row_buffer = remainder.strip()


def render_table(document: Document, buffer: TableBuffer):
    if buffer.raw_row_buffer.strip():
        add_table_row(buffer, buffer.raw_row_buffer)
        buffer.raw_row_buffer = ""

    if not buffer.rows:
        add_caption_paragraph(document, "Table", buffer.caption)
        return

    max_columns = max(len(row) for row in buffer.rows)
    table = document.add_table(rows=len(buffer.rows), cols=max_columns)
    table.style = "Table Grid"

    for row_index, row in enumerate(buffer.rows):
        for col_index in range(max_columns):
            value = row[col_index] if col_index < len(row) else ""
            table.cell(row_index, col_index).text = value

    add_caption_paragraph(document, "Table", buffer.caption)


def resolve_graphics_path(raw_path: str) -> Path | None:
    candidate = (ROOT / latex_to_text(raw_path)).resolve()
    if candidate.exists():
        return candidate

    for suffix in (".png", ".jpg", ".jpeg", ".bmp", ".gif"):
        with_suffix = candidate.with_suffix(suffix)
        if with_suffix.exists():
            return with_suffix
    return None


def build_docx(output_path: Path, google_docs_mode: bool = False):
    tex = MAIN_TEX.read_text(encoding="utf-8", errors="ignore")
    lines = tex.splitlines()
    table_captions, figure_captions = collect_caption_metadata(lines)
    outline = collect_outline_metadata(lines)

    document = Document()
    add_front_matter(document)

    dissertation_title = extract_macro(tex, "title", "Dissertation")
    author_name = extract_macro(tex, "author", "Author")

    title_lines = [
        "Glasgow Caledonian University",
        "School of Computing, Engineering and Built Environment",
        "MSc Computer Science Dissertation",
        "Agentic AI for Academic Planning",
        dissertation_title,
        f"Author: {author_name}",
        "Student Number: S2461801",
        "Degree: Master of Science in Computer Science",
        "Submission: March 2026",
        "",
        "A dissertation submitted in partial fulfilment of the requirements of Glasgow Caledonian University for the degree of Master of Science in Computer Science.",
        "This project report is my own original work and has not been submitted elsewhere in fulfilment of the requirements of this or any other award.",
    ]
    for text in title_lines:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = 1

    document.add_section(WD_SECTION_START.NEW_PAGE)

    in_document = False
    in_verbatim = False
    in_titlepage = False
    in_itemize = False
    in_enumerate = False
    in_figure = False
    in_table = False
    in_tabular = False
    body_section_started = False
    heading_counters = HeadingCounters()
    heading_numbering_active = False
    figure_buffer: FigureBuffer | None = None
    pending_paragraph_lines: list[str] = []
    table_buffer: TableBuffer | None = None
    enumerate_stack: list[int] = []

    for raw in lines:
        line = raw.rstrip()

        if r"\begin{document}" in line:
            in_document = True
            continue
        if not in_document:
            continue
        if r"\end{document}" in line:
            break

        if r"\begin{verbatim}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            in_verbatim = True
            continue
        if r"\end{verbatim}" in line:
            in_verbatim = False
            continue
        if in_verbatim:
            document.add_paragraph(line)
            continue

        if r"\begin{titlepage}" in line:
            in_titlepage = True
            continue
        if r"\end{titlepage}" in line:
            in_titlepage = False
            continue
        if in_titlepage:
            continue

        if r"\begin{figure}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            in_figure = True
            figure_buffer = FigureBuffer()
            continue
        if in_figure:
            image_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}", line)
            if image_match and figure_buffer is not None:
                figure_buffer.image_path = resolve_graphics_path(image_match.group(1))
            caption_match = re.search(r"\\caption\{(.+?)\}", line)
            if r"\end{figure}" in line:
                if caption_match and figure_buffer is not None:
                    figure_buffer.caption = clean_inline(caption_match.group(1))
                if figure_buffer and figure_buffer.image_path and figure_buffer.image_path.exists():
                    document.add_picture(str(figure_buffer.image_path), width=Mm(160))
                if figure_buffer:
                    add_caption_paragraph(document, "Figure", figure_buffer.caption)
                in_figure = False
                figure_buffer = None
            elif caption_match and figure_buffer is not None:
                figure_buffer.caption = clean_inline(caption_match.group(1))
            continue

        if r"\begin{table}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            in_table = True
            table_buffer = TableBuffer()
            continue
        if r"\end{table}" in line:
            if table_buffer:
                render_table(document, table_buffer)
            in_table = False
            in_tabular = False
            table_buffer = None
            continue

        if r"\begin{tabular}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            if table_buffer is None:
                table_buffer = TableBuffer()
            in_tabular = True
            continue
        if r"\end{tabular}" in line:
            in_tabular = False
            if not in_table and table_buffer:
                render_table(document, table_buffer)
                table_buffer = None
            continue

        if in_tabular and table_buffer is not None:
            append_table_source(table_buffer, line)
            continue

        if in_table and table_buffer is not None:
            caption_match = re.search(r"\\caption\{(.+?)\}", line)
            if caption_match:
                table_buffer.caption = clean_inline(caption_match.group(1))
            continue

        if r"\newpage" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            document.add_page_break()
            continue

        if r"\begin{itemize}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            in_itemize = True
            continue
        if r"\end{itemize}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            in_itemize = False
            continue

        if r"\begin{enumerate}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            enumerate_stack.append(0)
            in_enumerate = True
            continue
        if r"\end{enumerate}" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            if enumerate_stack:
                enumerate_stack.pop()
            in_enumerate = bool(enumerate_stack)
            continue

        section = re.match(r"\\section(\*?)\{(.+?)\}", line)
        if section:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            starred = bool(section.group(1))
            title = clean_inline(section.group(2))
            if title == "Introduction" and not body_section_started:
                document.add_section(WD_SECTION_START.NEW_PAGE)
                body_section_started = True
            display_title = title
            if not starred:
                heading_counters.section += 1
                heading_counters.subsection = 0
                heading_counters.subsubsection = 0
                heading_numbering_active = True
                display_title = f"{heading_counters.section} {title}"
            document.add_heading(display_title, level=1)
            continue

        subsection = re.match(r"\\subsection(\*?)\{(.+?)\}", line)
        if subsection:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            starred = bool(subsection.group(1))
            title = clean_inline(subsection.group(2))
            display_title = title
            if not starred and heading_numbering_active:
                heading_counters.subsection += 1
                heading_counters.subsubsection = 0
                display_title = f"{heading_counters.section}.{heading_counters.subsection} {title}"
            document.add_heading(display_title, level=2)
            continue

        subsubsection = re.match(r"\\subsubsection(\*?)\{(.+?)\}", line)
        if subsubsection:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            starred = bool(subsubsection.group(1))
            title = clean_inline(subsubsection.group(2))
            display_title = title
            if not starred and heading_numbering_active:
                heading_counters.subsubsection += 1
                display_title = f"{heading_counters.section}.{heading_counters.subsection}.{heading_counters.subsubsection} {title}"
            document.add_heading(display_title, level=3)
            continue

        if r"\tableofcontents" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            document.add_heading("Table of Contents", level=1)
            if google_docs_mode:
                add_static_toc(document, outline)
            else:
                add_toc(document)
            continue

        if r"\listoffigures" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            document.add_heading("List of Figures", level=1)
            if google_docs_mode:
                add_static_caption_list(document, "Figure", figure_captions)
            else:
                add_caption_index(document, "Figure")
            continue

        if r"\listoftables" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            document.add_heading("List of Tables", level=1)
            if google_docs_mode:
                add_static_caption_list(document, "Table", table_captions)
            else:
                add_caption_index(document, "Table")
            continue

        if r"\printbibliography" in line:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            document.add_heading("References", level=1)
            for reference in parse_bib_entries(REF_BIB):
                document.add_paragraph(reference, style="List Bullet")
            continue

        item = re.match(r"\s*\\item\s*(.*)", line)
        if item:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            item_text = clean_inline(item.group(1))
            if item_text:
                if in_enumerate:
                    if enumerate_stack:
                        enumerate_stack[-1] += 1
                    add_static_numbered_item(document, len(enumerate_stack), enumerate_stack[-1], item_text)
                else:
                    document.add_paragraph(item_text, style="List Bullet")
            continue

        if line.strip().startswith("\\"):
            continue

        cleaned = clean_inline(line)
        if not cleaned:
            pending_paragraph_lines = flush_pending_paragraph(document, pending_paragraph_lines)
            continue
        pending_paragraph_lines.append(cleaned)

    flush_pending_paragraph(document, pending_paragraph_lines)
    if not google_docs_mode:
        apply_page_numbering(document)
    document.save(str(output_path))


if __name__ == "__main__":
    build_docx(OUTPUT, google_docs_mode=False)
    build_docx(GOOGLE_OUTPUT, google_docs_mode=True)
