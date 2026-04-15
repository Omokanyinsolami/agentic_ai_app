from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


DOCX_PATH = Path("agentic ai dissertation.docx")


def backup_file(path: Path) -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = path.with_name(f"{path.stem}.backup_before_citation_expansion_{stamp}{path.suffix}")
    shutil.copy2(path, backup)
    return backup


def append_to_matching_paragraph(doc: Document, prefix: str, addition: str) -> bool:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(prefix):
            if addition not in text:
                paragraph.text = paragraph.text.rstrip() + " " + addition
            return True
    return False


def replace_reference(doc: Document, startswith: str, replacement: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            paragraph.text = replacement
            return True
    return False


def citation_audit(doc: Document) -> tuple[int, int, list[str]]:
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    try:
        ref_index = paragraphs.index("References")
    except ValueError:
        ref_index = len(paragraphs)

    body = "\n".join(p for p in paragraphs[:ref_index] if p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        body += "\n" + text

    citation_matches = re.findall(
        r"\(([A-Z][A-Za-zÀ-ÖØ-öø-ÿ .&,-]+?),\s*(20\d{2}|19\d{2})\)",
        body,
    )
    unique_citations = {f"{author.strip()}, {year}" for author, year in citation_matches}

    refs: list[tuple[str, str, str]] = []
    for text in paragraphs[ref_index + 1 :]:
        if not text or text.startswith("Appendix") or re.match(r"^\d+(\.\d+)?\s+", text):
            break
        match = re.match(r"^(.+?)\s*\((\d{4})\)", text)
        if match:
            refs.append((match.group(1).strip(), match.group(2), text))

    def expected_citation(author_field: str, year: str) -> str:
        if author_field in {
            "CrewAI",
            "LangChain AI",
            "Meta Platforms",
            "The PostgreSQL Global Development Group",
        }:
            return f"{author_field}, {year}"

        if " et al." in author_field:
            first = author_field.split(" et al.")[0].split(",")[0].strip()
            return f"{first} et al., {year}"

        # Corporate/online entries without initials.
        if "," not in author_field:
            return f"{author_field}, {year}"

        names = [part.strip() for part in re.split(r"\s+&\s+|\s+and\s+", author_field)]
        surnames = [name.split(",")[0].strip() for name in names if name.strip()]
        if len(surnames) == 1:
            return f"{surnames[0]}, {year}"
        if len(surnames) == 2:
            return f"{surnames[0]} and {surnames[1]}, {year}"
        return f"{surnames[0]} et al., {year}"

    missing: list[str] = []
    for author_field, year, _ in refs:
        expected = expected_citation(author_field, year)
        if expected not in unique_citations:
            missing.append(expected)

    return len(citation_matches), len(unique_citations), missing


def main() -> None:
    backup = backup_file(DOCX_PATH)
    doc = Document(DOCX_PATH)

    additions = [
        (
            "Learners increasingly manage academic work across fragmented digital environments",
            "Educational data-mining research also shows that structured learning activity data can support better understanding of learner progress when it is used meaningfully (Koedinger et al., 2010).",
        ),
        (
            "The project is motivated by a practical support problem and a technical design problem.",
            "This concern is consistent with wider discussions of agentic AI adoption, where capability must be balanced against reliability, governance, and user trust (Murugesan, 2025).",
        ),
        (
            "Three specific motivations follow from this.",
            "The shift from generative AI toward agentic AI also places stronger emphasis on action, feedback, and operational accountability rather than text generation alone (Schneider, 2025).",
        ),
        (
            "Alongside formal educational platforms, many learners also rely on calendars",
            "End-user planning research similarly shows that planning tools need to support user reasoning and plan correction rather than only store plan data (Lee et al., 2025).",
        ),
        (
            "What is often missing is a coherent supportive environment",
            "Automated planning research in dynamic domains also highlights the importance of connecting plan generation with changing situational constraints (Bidoux et al., 2017).",
        ),
        (
            "This project responds to that gap by treating learner support as an operational design objective.",
            "This is consistent with recent agentic AI surveys that position autonomous support as an applied systems problem rather than only a modelling problem (Pati, 2025).",
        ),
        (
            "Non-agentic AI systems in education usually provide bounded assistance",
            "Classical AI literature distinguishes such bounded task support from broader goal-directed agency and reasoning (Russell and Norvig, 2021).",
        ),
        (
            "Classical planning and optimization methods remain highly relevant in this space.",
            "Concise planning models further show why explicit state, action, and goal representations are useful for deterministic schedule construction (Geffner and Bonet, 2013).",
        ),
        (
            "Agentic AI extends beyond bounded recommendation by introducing goal-directed autonomy",
            "Conceptual taxonomies of AI agents and agentic AI also emphasise the distinction between isolated tool use and sustained goal-oriented behaviour (Sapkota et al., 2025).",
        ),
        (
            "This capability is especially relevant where academic work is dynamic.",
            "LLM-planning surveys show that language models can support planning, but they still require constraint management and evaluation against explicit task goals (Aghzal et al., 2025).",
        ),
        (
            "The software overview for this project focuses on frameworks",
            "Reasoning-and-acting patterns such as ReAct further motivate the combination of language reasoning with external action execution (Yao et al., 2022).",
        ),
        (
            "In addition to conceptual literature, the project required a technical review",
            "Research on self-evolving agents reinforces the need to examine architecture, feedback loops, and control boundaries rather than treating autonomy as a single feature (Fang et al., 2025).",
        ),
        (
            "Three architecture patterns were examined at a technical level.",
            "Surveys of LLM-based autonomous agents identify similar risks around autonomy, reliability, and tool-mediated execution (Chen et al., 2023).",
        ),
        (
            "The second pattern is a purely deterministic service architecture",
            "This baseline reflects the classical AI view that explicit representations and algorithmic control can support predictable reasoning, even when they are less flexible than adaptive agents (Russell and Norvig, 2021).",
        ),
        (
            "The technical review also considered orchestration granularity.",
            "Planning-with-search work also illustrates why reasoning depth and computational efficiency must be balanced in practical planning systems (Katz et al., 2024).",
        ),
        (
            "At the data layer, PostgreSQL was selected",
            "The selection also aligns with PostgreSQL's documented emphasis on relational integrity, indexing, and transactional database behaviour (The PostgreSQL Global Development Group, 2026).",
        ),
        (
            "The frontend technical review examined how interface architecture",
            "React was selected because its component model supports structured interface state, reusable views, and interactive user workflows (Meta Platforms, 2026).",
        ),
        (
            "This project is framed as a design-and-implementation study",
            "Recent research landscapes of agentic AI and LLMs also emphasise that evaluation should consider applications, limitations, and deployment implications together (Brohi et al., 2025).",
        ),
        (
            "At scheduling level, alternatives included pure heuristic scheduling",
            "Adaptive planning research also indicates that feedback can improve plan refinement when plans must be revised after environmental change (Sun et al., 2023).",
        ),
        (
            "Adaptation trigger strategy was another explicit choice.",
            "Agent theories that connect goals, actions, and changing environments support this focus on meaningful state transitions rather than constant recomputation (Novák and Jamroga, 2011).",
        ),
        (
            "Compared to pure LLM-based agents",
            "This trade-off is also reflected in surveys of agentic AI architectures, where autonomy, flexibility, evaluation, and operational control remain active design tensions (Bandi et al., 2025).",
        ),
        (
            "Compared to existing end-user planning research prototypes",
            "End-user planning work is therefore useful as a comparison point because it highlights plan verification and inspectability as practical user-facing concerns (Lee et al., 2025).",
        ),
        (
            "The implemented architecture demonstrates a pattern",
            "Current agentic AI governance guidance supports this pattern because bounded autonomy and explicit control points make systems easier to inspect and manage (Shavit et al., 2023).",
        ),
        (
            "The discussion has direct implications",
            "This is consistent with multi-expert analysis of agentic systems, which stresses reliability, human oversight, and clear operational boundaries in user-facing applications (Hughes et al., 2025).",
        ),
        (
            "Although the project demonstrates robust behaviour",
            "This boundary is important because agentic systems are increasingly discussed across many domains, but transfer depends on the fit between task structure and system capability (Fang et al., 2025).",
        ),
        (
            "Methodological rigour in software dissertations",
            "This evaluation stance is consistent with planning literature, where explicit metrics and reproducible scenarios are needed to compare planning approaches defensibly (Ghallab et al., 2016).",
        ),
        (
            "This project has demonstrated that agentic AI principles",
            "The selected LangGraph approach is also consistent with framework documentation that emphasises controllable, stateful orchestration for agent workflows (LangChain AI, 2024).",
        ),
        (
            "The hybrid approach of combining deterministic graph-based routing",
            "This balanced interpretation follows recent agentic AI literature, which treats autonomous capability and reliability constraints as connected design concerns (Sapkota et al., 2025).",
        ),
        (
            "This research project contributes to understanding",
            "The contribution therefore sits within the wider movement from generative interaction toward agentic systems that can act, monitor, and revise behaviour over time (Schneider, 2025).",
        ),
    ]

    missing_targets = []
    for prefix, addition in additions:
        if not append_to_matching_paragraph(doc, prefix, addition):
            missing_targets.append(prefix)

    references = {
        "Acharya, D. B.": "Acharya, D. B., Kuppan, K. and Divya, B. (2025) 'Agentic AI: Autonomous Intelligence for Complex Goals - A Comprehensive Survey', IEEE Access, 13, pp. 18912-18936. doi: 10.1109/ACCESS.2025.3532853.",
        "Aghzal, M.": "Aghzal, M., Plaku, E., Stein, G. J. and Yao, Z. (2025) 'A survey on large language models for automated planning', arXiv preprint arXiv:2502.12435.",
        "Bandi, A.": "Bandi, A., Kongari, B., Naguru, R., Pasnoor, S. and Vilipala, S. V. (2025) 'The rise of agentic AI: A review of definitions, frameworks, architectures, applications, evaluation metrics, and challenges', Future Internet, 17(9), article 404.",
        "Bidoux, L.": "Bidoux, L., Pignon, J. P. and Bénaben, F. (2017) 'On the use of automated planning for crisis management', Proceedings of the International Conference on Information Systems for Crisis Response and Management (ISCRAM).",
        "Brohi, S.": "Brohi, S., Jhanjhi, N. and Pillai, T. R. (2025) 'A research landscape of agentic AI and large language models: Applications, challenges and future directions', Algorithms, 18(8), article 499.",
        "Chase, H.": "Chase, H. (2022) LangChain [online]. Available at: https://github.com/langchain-ai/langchain (Accessed: 9 April 2026).",
        "Chen, Z.": "Chen, Z., Tang, J., Chen, X. and Lin, Y. (2023) 'A survey on large language model based autonomous agents', arXiv preprint.",
        "CrewAI": "CrewAI (2024) CrewAI [online]. Available at: https://github.com/crewAIInc/crewAI (Accessed: 9 April 2026).",
        "Fang, J.": "Fang, J. et al. (2025) 'A comprehensive survey of self-evolving AI agents: A new paradigm bridging foundation models and lifelong agentic systems', arXiv preprint arXiv:2508.07407.",
        "Geffner, H.": "Geffner, H. and Bonet, B. (2013) A Concise Introduction to Models and Methods for Automated Planning. Cham: Springer.",
        "Ghallab, M.": "Ghallab, M., Nau, D. and Traverso, P. (2016) Automated Planning and Acting. Cambridge: Cambridge University Press.",
        "Hughes, L.": "Hughes, L. et al. (2025) 'AI agents and agentic systems: A multi-expert analysis', Journal of Computer Information Systems, pp. 1-29.",
        "Russell, S.": "Russell, S. and Norvig, P. (2021) Artificial Intelligence: A Modern Approach. 4th edn. Harlow: Pearson.",
        "Katz, M.": "Katz, M., Kokel, H., Srinivas, K. and Sohrabi Araghi, S. (2024) 'Thought of search: Planning with language models through the lens of efficiency', Advances in Neural Information Processing Systems, 37, pp. 138491-138568.",
        "Koedinger, K. R.": "Koedinger, K. R., Baker, R. S., Cunningham, K., Skogsholm, A., Leber, B. and Stamper, J. (2010) 'A data repository for the EDM community: The PSLC DataShop', in Romero, C., Ventura, S., Pechenizkiy, M. and Baker, R. S. J. d. (eds.) Handbook of Educational Data Mining. Boca Raton, FL: CRC Press, pp. 43-56.",
        "LangChain AI": "LangChain AI (2024) LangGraph overview [online]. Available at: https://docs.langchain.com/oss/python/langgraph/overview (Accessed: 9 April 2026).",
        "Lee, C. P.": "Lee, C. P., Porfirio, D., Wang, X. J., Zhao, K. C. and Mutlu, B. (2025) 'VeriPlan: Integrating formal verification and LLMs into end-user planning', Proceedings of the 2025 CHI Conference on Human Factors in Computing Systems, pp. 1-19.",
        "Meta Platforms": "Meta Platforms (2026) React documentation [online]. Available at: https://react.dev/ (Accessed: 9 April 2026).",
        "Murugesan, S.": "Murugesan, S. (2025) 'The rise of agentic AI: Implications, concerns, and the path forward', IEEE Intelligent Systems, 40(2), pp. 8-14.",
        "Novák, P.": "Novák, P. and Jamroga, W. (2011) 'Agents, actions and goals in dynamic environments', Proceedings of the International Joint Conference on Artificial Intelligence (IJCAI), pp. 313-318.",
        "Pati, A. K.": "Pati, A. K. (2025) 'Agentic AI: A comprehensive survey of technologies, applications, and societal implications', IEEE Access.",
        "Ronacher, A.": "Ronacher, A. (2010) Flask [online]. Available at: https://flask.palletsprojects.com/ (Accessed: 9 April 2026).",
        "Sapkota, R.": "Sapkota, R., Roumeliotis, K. I. and Karkee, M. (2025) 'AI agents vs. agentic AI: A conceptual taxonomy, applications and challenges', arXiv preprint arXiv:2505.10468.",
        "Schneider, J.": "Schneider, J. (2025) 'Generative to agentic AI: Survey, conceptualization, and challenges', arXiv preprint arXiv:2504.18875.",
        "Shavit, Y.": "Shavit, Y., Agarwal, S., Brundage, M. et al. (2023) Practices for Governing Agentic AI Systems [online]. OpenAI. Available at: https://openai.com/index/practices-for-governing-agentic-ai-systems/ (Accessed: 9 April 2026).",
        "Sun, H.": "Sun, H., Zhuang, Y., Kong, L., Dai, B. and Zhang, C. (2023) 'AdaPlanner: Adaptive planning from feedback with language models', Advances in Neural Information Processing Systems, 36, pp. 58202-58245.",
        "The PostgreSQL Global Development Group": "The PostgreSQL Global Development Group (2026) PostgreSQL documentation [online]. Available at: https://www.postgresql.org/docs/ (Accessed: 9 April 2026).",
        "Wang, L.": "Wang, L. et al. (2024) 'A survey on large language model based autonomous agents', Frontiers of Computer Science, 18(6), article 186345.",
        "Wu, Q.": "Wu, Q. et al. (2023) 'AutoGen: Enabling next-generation LLM applications via multi-agent conversation' [online]. Available at: https://github.com/microsoft/autogen (Accessed: 9 April 2026).",
        "Yao, S.": "Yao, S. et al. (2022) 'ReAct: Synergizing reasoning and acting in language models', The Eleventh International Conference on Learning Representations.",
    }

    for prefix, replacement in references.items():
        replace_reference(doc, prefix, replacement)

    doc.save(DOCX_PATH)
    updated = Document(DOCX_PATH)
    citation_count, unique_count, missing = citation_audit(updated)
    print(f"backup={backup}")
    print(f"citation_count={citation_count}")
    print(f"unique_citation_count={unique_count}")
    print(f"missing_reference_citations={missing}")
    if missing_targets:
        print("missing_targets=" + repr(missing_targets))


if __name__ == "__main__":
    main()
